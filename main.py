import os
import sys
import io
import winsound
from datetime import datetime

import keyboard
from PIL import ImageGrab
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT


CAPTURE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "captures")
HOTKEY = "ctrl+shift+s"


def create_document():
    """Buat document Word baru dengan landscape orientation."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Swap width and height for landscape
    section.page_width, section.page_height = section.page_height, section.page_width
    # Kurangi margin agar gambar lebih besar
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    return doc


def capture_screenshot(doc, counter, filepath):
    """Ambil screenshot, tambahkan ke document, save, dan play beep."""
    try:
        screenshot = ImageGrab.grab()

        # Simpan screenshot ke BytesIO (tanpa kompresi)
        img_buffer = io.BytesIO()
        screenshot.save(img_buffer, format="PNG")
        img_buffer.seek(0)

        # Hitung lebar yang tersedia di halaman
        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin

        # Tambahkan gambar ke document dengan lebar penuh halaman
        doc.add_picture(img_buffer, width=available_width)

        # Tambahkan caption
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        caption = doc.add_paragraph(f"Screenshot #{counter} — {timestamp}")
        caption.style = doc.styles["Normal"]
        caption_format = caption.paragraph_format
        caption_format.space_after = Inches(0.3)

        # Save document setiap kali capture (supaya tidak hilang kalau crash)
        doc.save(filepath)

        # Feedback: beep sound
        winsound.Beep(1000, 150)

        print(f"  Screenshot #{counter} captured at {timestamp}")

    except Exception as e:
        print(f"  Error capturing screenshot: {e}")
        # Error beep
        winsound.Beep(400, 300)


def main():
    os.makedirs(CAPTURE_DIR, exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    filename = f"capture_{timestamp}.docx"
    filepath = os.path.join(CAPTURE_DIR, filename)

    doc = create_document()
    counter = 0

    print("=" * 50)
    print("  Auto Screen Capture")
    print("=" * 50)
    print(f"  Output : {filepath}")
    print(f"  Capture: {HOTKEY.upper()}")
    print(f"  Quit   : Ctrl+C")
    print("=" * 50)
    print()

    def on_hotkey():
        nonlocal counter
        counter += 1
        capture_screenshot(doc, counter, filepath)

    keyboard.add_hotkey(HOTKEY, on_hotkey, suppress=True)

    try:
        keyboard.wait()
    except KeyboardInterrupt:
        pass

    print()
    print("=" * 50)
    print(f"  Session ended. Total screenshots: {counter}")
    if counter > 0:
        print(f"  Saved to: {filepath}")
    else:
        # Hapus file kosong jika tidak ada screenshot
        if os.path.exists(filepath):
            os.remove(filepath)
        print("  No screenshots taken. File not created.")
    print("=" * 50)


if __name__ == "__main__":
    main()
